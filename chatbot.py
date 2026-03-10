"""
Thesis Chatbot - Full-Document Prompt Engine
Reads the entire thesis .docx and sends it as context with each OpenAI API call.
No RAG, no embeddings, no vector database — just clean, simple prompting.
"""

import os
import re
import requests
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables
load_dotenv(override=True)

# --- Configuration ---
OPENAI_MODEL = "gpt-4o-mini"
OPENAI_URL = "https://api.openai.com/v1/chat/completions"

# Paths
DATA_DIR = Path(__file__).parent / "data"

# In-memory cache of the thesis text
_thesis_text = None
_is_loaded = False


def _get_openai_api_key() -> str:
    """
    Resolve API key at call time so updates to .env are picked up
    without requiring a process restart.
    """
    load_dotenv(override=True)
    return os.getenv("OPENAI_API_KEY", "").strip().strip('"').strip("'")


def _load_docx(file_path: str) -> str:
    """Extract all text from a .docx file (paragraphs + tables)."""
    from docx import Document
    doc = Document(file_path)
    full_text = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            full_text.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    return "\n".join(full_text)


def load_thesis():
    """
    Load the thesis document into memory.
    Called once at startup. The full text is cached for all future requests.
    """
    global _thesis_text, _is_loaded

    if _is_loaded and _thesis_text:
        print("[OK] Thesis document already loaded.")
        return True

    try:
        thesis_files = list(DATA_DIR.glob("*.docx"))
        if not thesis_files:
            print("[ERROR] No .docx thesis file found in data/ directory.")
            return False

        thesis_path = thesis_files[0]
        print(f"[INFO] Loading thesis document: {thesis_path.name}")

        _thesis_text = _load_docx(str(thesis_path))
        if not _thesis_text:
            print("[ERROR] No text extracted from thesis document.")
            return False

        _is_loaded = True
        print(f"[OK] Thesis loaded — {len(_thesis_text):,} characters in memory.")
        return True

    except ImportError:
        print("[ERROR] Missing dependency: python-docx")
        print("   Run: pip install python-docx")
        return False
    except Exception as e:
        print(f"[ERROR] Failed to load thesis: {e}")
        return False


def ask_thesis(question: str) -> dict:
    """
    Answer a question about the thesis.
    Sends the full document text + question to OpenAI and returns the response.
    """
    global _is_loaded, _thesis_text

    # Ensure thesis is loaded
    if not _is_loaded:
        success = load_thesis()
        if not success:
            return {
                "answer": "I'm sorry, the thesis document hasn't been loaded yet. "
                          "Please make sure the .docx file is in the data/ folder.",
                "sources": []
            }

    # Check API key
    openai_api_key = _get_openai_api_key()
    if not openai_api_key or openai_api_key == "your-api-key-here":
        return {
            "answer": "The OpenAI API key hasn't been configured yet. "
                      "Please add your API key to the .env file.",
            "sources": []
        }

    # Build the system prompt with the full thesis text
    system_prompt = (
        'You are a helpful research assistant for a thesis document titled '
        '"Enhancing Localized Weather Forecasting in Iloilo Using '
        'Long-Short Term Memory (LSTM)".\n\n'
        "STRICT RULES:\n"
        "1. Answer ONLY based on the research paper text provided below. "
        "Never invent or assume details.\n"
        "2. If the answer is not found in the document, reply exactly: "
        '"That information is not explicitly stated in the research paper."\n'
        "3. If the question is unrelated to the research paper or thesis project, "
        "reply exactly: "
        '"I can only answer questions related to the research paper and this thesis project."\n'
        "4. Be concise but thorough. Use bullet points when listing multiple items.\n"
        "5. Reference specific findings, data, or sections when relevant.\n"
        "6. Maintain an academic but approachable tone.\n"
        "7. Do NOT use markdown headers (# or ##). Keep responses as flowing text "
        "with bullet points where needed.\n\n"
        "FULL RESEARCH PAPER TEXT:\n"
        "─────────────────────────\n"
        f"{_thesis_text}\n"
        "─────────────────────────"
    )

    # Call OpenAI API
    try:
        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": OPENAI_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question}
            ],
            "max_tokens": 1024,
            "temperature": 0.3
        }

        response = requests.post(
            OPENAI_URL,
            headers=headers,
            json=payload,
            timeout=90
        )

        if response.status_code == 200:
            data = response.json()
            answer = data["choices"][0]["message"]["content"]

            # Clean up any thinking tags if present
            answer = re.sub(r'<think>.*?</think>', '', answer, flags=re.DOTALL).strip()

            return {
                "answer": answer,
                "sources": []
            }
        else:
            error_msg = response.text
            print(f"[ERROR] OpenAI API error ({response.status_code}): {error_msg}")

            if response.status_code == 401:
                friendly = ("OpenAI rejected the API key (401 Unauthorized). "
                            "Please verify OPENAI_API_KEY in .env.")
            elif response.status_code == 429:
                friendly = ("OpenAI rate limit or quota issue (429). "
                            "Check your billing and usage limits.")
            else:
                friendly = (f"I encountered an error communicating with the AI service. "
                            f"Please try again later. (Status: {response.status_code})")

            return {"answer": friendly, "sources": []}

    except requests.exceptions.Timeout:
        return {
            "answer": "The AI service took too long to respond. Please try again.",
            "sources": []
        }
    except Exception as e:
        print(f"[ERROR] Error in ask_thesis: {e}")
        return {
            "answer": "An unexpected error occurred. Please try again later.",
            "sources": []
        }


def get_status() -> dict:
    """Return the current status of the chatbot."""
    openai_api_key = _get_openai_api_key()
    return {
        "initialized": _is_loaded,
        "api_key_set": bool(openai_api_key and openai_api_key != "your-api-key-here"),
        "model": OPENAI_MODEL,
        "data_dir": str(DATA_DIR),
        "thesis_files": [f.name for f in DATA_DIR.glob("*.docx")] if DATA_DIR.exists() else []
    }
