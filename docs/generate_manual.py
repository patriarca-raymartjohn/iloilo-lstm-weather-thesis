# -*- coding: utf-8 -*-
"""
Generate the User Manual (Word .docx) for the
Iloilo Weather - LSTM Dashboard thesis application.
Screenshots are read from docs/manual_assets/.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "manual_assets")
OUT_DOCX = os.path.join(HERE, "Iloilo_Weather_LSTM_User_Manual.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # deep blue
ACCENT2 = RGBColor(0x2E, 0x74, 0xB5)     # lighter blue
GREY = RGBColor(0x59, 0x59, 0x59)

DEPLOYED_URL = "https://iloilo-lstm-weather-thesis.onrender.com/"


def img(name):
    return os.path.join(ASSETS, name)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    # PAGE field
    fldStart = OxmlElement('w:fldSimple')
    fldStart.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    r.append(rpr)
    t = OxmlElement('w:t')
    t.text = "1"
    r.append(t)
    fldStart.append(r)
    p._p.append(fldStart)
    run2 = p.add_run(" of ")
    run2.font.size = Pt(9)
    run2.font.color.rgb = GREY
    fldPages = OxmlElement('w:fldSimple')
    fldPages.set(qn('w:instr'), 'NUMPAGES')
    r2 = OxmlElement('w:r')
    r2.append(OxmlElement('w:rPr'))
    t2 = OxmlElement('w:t')
    t2.text = "1"
    r2.append(t2)
    fldPages.append(r2)
    p._p.append(fldPages)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level == 1 else ACCENT2
    return h


def body(doc, text, size=11, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(doc, text, bold_lead=None):
    p = doc.add_paragraph(style='List Number')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F3F5')
    pPr.append(shd)
    return p


def figure(doc, filename, caption, width=6.2):
    if not os.path.exists(img(filename)):
        body(doc, f"[Screenshot missing: {filename}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img(filename), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(12)


def page_break(doc):
    doc.add_page_break()


# ----------------------------------------------------------------------------
doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

section = doc.sections[0]
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# ============================= COVER PAGE ===================================
for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("USER MANUAL")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Iloilo Weather – LSTM Forecasting Dashboard")
r.bold = True
r.font.size = Pt(17)
r.font.color.rgb = ACCENT2

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A 2-Day Localized Weather Forecasting Web Application")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = GREY

doc.add_paragraph()
# cover image
if os.path.exists(img("01_dashboard_top.png")):
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run().add_picture(img("01_dashboard_top.png"), width=Inches(5.6))

doc.add_paragraph()
def center(text, size=12, bold=False, color=None, italic=False):
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(text)
    rr.bold = bold
    rr.italic = italic
    rr.font.size = Pt(size)
    if color:
        rr.font.color.rgb = color
    pp.paragraph_format.space_after = Pt(2)
    return pp

center("Based on the undergraduate thesis", size=11, italic=True, color=GREY)
center("Enhancing Localized Weather Forecasting in Iloilo", size=13, bold=True, color=ACCENT)
center("Using Long Short-Term Memory (LSTM)", size=13, bold=True, color=ACCENT)
doc.add_paragraph()
center("Researchers", size=11, bold=True, color=ACCENT2)
center("Nethan Quinn G. Jael  •  Loi Marie Maxino", size=11)
center("Aaron Hans L. Oliverio  •  Raymart John P. Patriarca", size=11)
doc.add_paragraph()
center("Bachelor of Science in Computer Science", size=11, bold=True)
center("Iloilo Science and Technology University", size=11)
center("La Paz, Iloilo City • Western Visayas, Philippines", size=10, color=GREY)
doc.add_paragraph()
center("March 2026", size=11, bold=True, color=GREY)

page_break(doc)

# ============================= TABLE OF CONTENTS ============================
heading(doc, "Table of Contents", level=1)
body(doc, "The entries below are linked to a live Word table of contents. "
          "In Microsoft Word, right-click the table and choose “Update Field” "
          "to refresh page numbers at any time.", size=10, italic=True)

# Real Word TOC field
p = doc.add_paragraph()
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fld_sep = OxmlElement('w:fldChar')
fld_sep.set(qn('w:fldCharType'), 'separate')
fld_text = OxmlElement('w:r')
ft = OxmlElement('w:t')
ft.text = "Right-click and select 'Update Field' to generate the table of contents."
fld_text.append(ft)
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
run = p.add_run()
run._r.append(fld_begin)
run._r.append(instr)
run._r.append(fld_sep)
run._r.append(fld_text)
run._r.append(fld_end)

page_break(doc)

# ============================= 1. INTRODUCTION =============================
heading(doc, "1. Introduction", level=1)
body(doc, "The Iloilo Weather – LSTM Dashboard is a web application that produces a "
          "direct 2-day weather forecast for Iloilo City and nearby Western Visayas "
          "locations. It is the software output of the thesis “Enhancing Localized "
          "Weather Forecasting in Iloilo Using Long Short-Term Memory (LSTM).”")
body(doc, "Unlike general weather services, the system is trained specifically on "
          "Iloilo’s historical climate. It takes the last 7 days of observed weather "
          "as input and predicts the next two days in a single model pass (a "
          "“direct multi-step” forecast). The dashboard also lets users compare the "
          "forecast against the past seven observed days and against the same dates "
          "in previous years, and includes an AI assistant that answers questions "
          "about the underlying research.")

heading(doc, "1.1 Who This Manual Is For", level=2)
body(doc, "This manual serves two audiences:")
bullet(doc, "who simply want to view forecasts and explore the dashboard "
            "on the deployed website.", bold_lead="End users – ")
bullet(doc, "instructors, panelists, or developers who want to run the "
            "application on their own computer.", bold_lead="Technical users – ")

heading(doc, "1.2 Key Features", level=2)
bullet(doc, "2-day temperature, rainfall, wind, humidity, dew point and sunshine forecast.", bold_lead="Direct LSTM forecast – ")
bullet(doc, "search any city/locality in the Iloilo area to re-center the forecast.", bold_lead="Location search – ")
bullet(doc, "interactive chart and table of the last 7 observed days.", bold_lead="Recent trends – ")
bullet(doc, "same-date weather across 2014–2025 for context.", bold_lead="Historical comparison – ")
bullet(doc, "export any data table as a CSV file.", bold_lead="CSV download – ")
bullet(doc, "ask questions about the thesis and get instant answers.", bold_lead="AI thesis assistant – ")
bullet(doc, "toggle a comfortable viewing theme.", bold_lead="Light / dark mode – ")

page_break(doc)

# ============================= 2. SYSTEM OVERVIEW ==========================
heading(doc, "2. How the System Works", level=1)
body(doc, "Understanding the data flow helps interpret what you see on screen:")
numbered(doc, "the app retrieves the last seven days of daily "
              "weather for the selected location from the Open-Meteo API.",
              bold_lead="Collect input – ")
numbered(doc, "the data is cleaned and scaled exactly the way it "
              "was during model training (Min-Max normalization).",
              bold_lead="Preprocess – ")
numbered(doc, "the trained LSTM model (input shape 7×11, output "
              "shape 2×7) predicts both upcoming days at once.",
              bold_lead="Predict – ")
numbered(doc, "predicted values are converted back to real-world "
              "units and shown as Day 1 and Day 2 forecast cards.",
              bold_lead="Display – ")
body(doc, "Forecasted variables: minimum and maximum temperature, rainfall amount, "
          "wind speed, relative humidity, dew point, and sunshine duration.")

t = doc.add_table(rows=1, cols=2)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].paragraphs[0].add_run("Component").bold = True
hdr[1].paragraphs[0].add_run("Technology").bold = True
rows = [
    ("Web framework", "Flask (Python 3.11)"),
    ("Forecast model", "Stacked LSTM neural network (Keras / TensorFlow)"),
    ("Weather data source", "Open-Meteo historical & forecast API"),
    ("AI thesis assistant", "OpenAI GPT-4o-mini (full-document prompting)"),
    ("Hosting (deployed)", "Render cloud platform"),
]
for a, b in rows:
    cells = t.add_row().cells
    cells[0].paragraphs[0].add_run(a)
    cells[1].paragraphs[0].add_run(b)
doc.add_paragraph()

page_break(doc)

# ============================= 3. DEPLOYED VERSION =========================
heading(doc, "3. Using the Deployed Website", level=1)
body(doc, "The fastest way to use the application is through the public website. "
          "No installation is required — only a web browser and an internet connection.")

heading(doc, "3.1 Opening the App", level=2)
numbered(doc, "Open any modern web browser (Chrome, Edge, Firefox, or Safari).")
p = doc.add_paragraph(style='List Number')
p.add_run("Go to the application address:")
code_block(doc, DEPLOYED_URL)
numbered(doc, "Wait for the dashboard to load. The default location is Iloilo City.")
note = doc.add_paragraph()
nr = note.add_run("Note:  The site is hosted on Render’s free tier. If it has been "
                  "idle, the first load may take 30–60 seconds while the server "
                  "“wakes up.” Subsequent loads are fast.")
nr.italic = True
nr.font.size = Pt(10)
nr.font.color.rgb = GREY

figure(doc, "01_dashboard_top.png",
       "Figure 3.1 — The dashboard header and the 2-day forecast cards (Day 1 and Day 2).")

heading(doc, "3.2 Reading the Forecast Cards", level=2)
body(doc, "The two cards at the top show the forecast for today/Day 1 and Day 2. The "
          "large number is the headline temperature; the table below each card lists "
          "the full set of predicted variables:")
bullet(doc, "the predicted low and high temperature in °C.", bold_lead="Min/Max Temp – ")
bullet(doc, "expected rainfall in millimetres.", bold_lead="Rain – ")
bullet(doc, "mean wind speed in km/h.", bold_lead="Wind – ")
bullet(doc, "average relative humidity (%).", bold_lead="Humid – ")
bullet(doc, "dew point temperature in °C.", bold_lead="Dew Pt – ")
bullet(doc, "total sunshine duration in seconds.", bold_lead="Sunshine – ")

page_break(doc)

heading(doc, "3.3 Searching for a Location", level=2)
body(doc, "By default the dashboard forecasts for Iloilo City. To view another locality:")
numbered(doc, "Click the search box in the header (“Search for location…”).")
numbered(doc, "Type a place name, for example “Jaro”.")
numbered(doc, "Choose the matching suggestion from the drop-down list.")
numbered(doc, "The dashboard reloads and recalculates the forecast for that location.")
figure(doc, "03_location_search.png",
       "Figure 3.2 — Typing in the location search box shows matching places with "
       "their coordinates and elevation.")
note = doc.add_paragraph()
nr = note.add_run("Tip:  A few small localities may not have a full seven days of recent "
                  "data available. If so, the app displays a message such as "
                  "“Need at least 7 days of history” — simply choose a larger nearby "
                  "city (e.g., Iloilo City) instead.")
nr.italic = True
nr.font.size = Pt(10)
nr.font.color.rgb = GREY

heading(doc, "3.4 Last 7 Observed Days — Trends", level=2)
body(doc, "This section plots the real weather of the past seven days that the model "
          "used as input. Use the coloured buttons (Min Temp, Max Temp, Rain, Wind, "
          "Humid, Dew Pt, Sunshine) to switch which variable the chart displays. The "
          "data table beneath the chart lists the exact values, and the "
          "“Download CSV” button exports them.")
figure(doc, "05_trends_section.png",
       "Figure 3.3 — The Last 7 Observed Days section with metric toggle buttons, "
       "chart, data table, and CSV export.")

page_break(doc)

heading(doc, "3.5 Historical Comparison", level=2)
body(doc, "To give the forecast context, this section shows what the weather was like "
          "on the same two calendar dates in previous years (2014–2025). Two panels "
          "— one for Day 1 and one for Day 2 — each provide the same metric toggle "
          "buttons and a per-year data table with its own CSV download.")
figure(doc, "06_historical_section.png",
       "Figure 3.4 — Historical comparison of the forecast dates across past years.")

heading(doc, "3.6 The AI Thesis Assistant", level=2)
body(doc, "A floating robot button in the bottom-right corner opens the Thesis "
          "Assistant. It can answer questions about the research paper behind the app.")
numbered(doc, "Click the robot (🤖) button to open the chat panel.")
numbered(doc, "Click one of the suggested questions, or type your own question.")
numbered(doc, "Press send and read the assistant’s reply.")
body(doc, "Example questions: “What is this study about?”, “What LSTM model was "
          "used?”, “What are the key findings?”, “Who are the researchers?”")
figure(doc, "07_chatbot.png",
       "Figure 3.5 — The Thesis Assistant answering “What is this study about?”")

heading(doc, "3.7 Light / Dark Mode", level=2)
body(doc, "Use the sun/moon toggle in the header to switch between light and dark "
          "themes for comfortable viewing in different lighting conditions.")

page_break(doc)

# ============================= 4. LOCAL VERSION ===========================
heading(doc, "4. Running the App Locally (Quick Start)", level=1)
body(doc, "Technical users can run the same application on their own computer. These "
          "quick-start steps assume the project files and Python environment are "
          "already present on the machine (as delivered with the thesis).")

heading(doc, "4.1 Requirements", level=2)
bullet(doc, "(the project is pinned to version 3.11).", bold_lead="Python 3.11 – ")
bullet(doc, "the project folder containing app.py, the model files, and the "
            ".venv virtual environment.", bold_lead="Project files – ")
bullet(doc, "an OpenAI API key in the .env file is required only for the AI "
            "assistant; the forecast works without it.", bold_lead="OpenAI key (optional) – ")

heading(doc, "4.2 Start the Application", level=2)
p = doc.add_paragraph(style='List Number')
p.add_run("Open a terminal (PowerShell) in the project folder and activate the "
          "virtual environment:")
code_block(doc, ".venv\\Scripts\\Activate.ps1")
p = doc.add_paragraph(style='List Number')
p.add_run("Start the server:")
code_block(doc, "python app.py")
p = doc.add_paragraph(style='List Number')
p.add_run("Wait for the startup messages. When ready, the terminal prints the local "
          "address it is serving on:")
code_block(doc,
           "[OK] Loaded model: LSTM_Weather_Forecast_Direct2_x7.keras\n"
           "[OK] Model input shape : (None, 7, 11)\n"
           "[OK] Model output shape: (None, 2, 7)\n"
           " * Running on http://127.0.0.1:5000")
note = doc.add_paragraph()
nr = note.add_run("By default the app serves on port 10000 (open "
                  "http://localhost:10000). To use a different port, set the PORT "
                  "environment variable before starting — for example PORT=5000 "
                  "serves on http://localhost:5000.")
nr.italic = True
nr.font.size = Pt(10)
nr.font.color.rgb = GREY

p = doc.add_paragraph(style='List Number')
p.add_run("Open the printed address in your browser. The dashboard looks and behaves "
          "exactly like the deployed site.")
figure(doc, "08_local_dashboard.png",
       "Figure 4.1 — The application running locally in the browser.")
note = doc.add_paragraph()
nr = note.add_run("Note:  The very first page load runs the model and downloads recent "
                  "weather data, so it can take up to a minute. Later requests are "
                  "much faster.")
nr.italic = True
nr.font.size = Pt(10)
nr.font.color.rgb = GREY

heading(doc, "4.3 Stopping the App", level=2)
body(doc, "Return to the terminal and press Ctrl + C to stop the local server.")

page_break(doc)

# ============================= 5. TROUBLESHOOTING =========================
heading(doc, "5. Troubleshooting", level=1)
t = doc.add_table(rows=1, cols=2)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].paragraphs[0].add_run("Problem").bold = True
hdr[1].paragraphs[0].add_run("What to do").bold = True
tb = [
    ("The deployed page takes a long time to load.",
     "The free Render server is waking from sleep. Wait 30–60 seconds and refresh once."),
    ("“Need at least 7 days of history” appears after searching a place.",
     "That locality lacks enough recent data. Choose a larger nearby city such as Iloilo City."),
    ("The forecast values look unchanged after switching location.",
     "Wait a moment — the forecast recalculates after the page reloads."),
    ("The AI assistant does not answer (local run).",
     "Ensure a valid OPENAI_API_KEY is set in the .env file, then restart the app."),
    ("Charts or tables are empty.",
     "Refresh the page; if it persists, check your internet connection (the app needs the Open-Meteo API)."),
    ("Local app: 'python is not recognized'.",
     "Activate the virtual environment first (.venv\\Scripts\\Activate.ps1)."),
]
for a, b in tb:
    cells = t.add_row().cells
    cells[0].paragraphs[0].add_run(a)
    cells[1].paragraphs[0].add_run(b)
doc.add_paragraph()

# ============================= 6. CREDITS =================================
heading(doc, "6. About & Credits", level=1)
body(doc, "Thesis title: Enhancing Localized Weather Forecasting in Iloilo Using Long "
          "Short-Term Memory (LSTM).")
body(doc, "Institution: Iloilo Science and Technology University — La Paz, Iloilo City.")
body(doc, "Degree: Bachelor of Science in Computer Science — March 2026.")
body(doc, "Researchers:", space_after=2)
bullet(doc, "Nethan Quinn G. Jael")
bullet(doc, "Loi Marie Maxino")
bullet(doc, "Aaron Hans L. Oliverio")
bullet(doc, "Raymart John P. Patriarca")
body(doc, "Adviser: Lovidrick Jhon P. Barrios, MSCS  •  Subject Professor: "
          "Maureen Nettie N. Linan, DIT.")
body(doc, "Panel: Joyce F. Jamile, MSCS  •  Yvette G. Gonzales, D. Eng.")
doc.add_paragraph()
body(doc, "Deployed application: " + DEPLOYED_URL, size=10, italic=True)

# Footer page numbers
add_page_number_footer(doc.sections[0])

doc.save(OUT_DOCX)
print("SAVED:", OUT_DOCX)
